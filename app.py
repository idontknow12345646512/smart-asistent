import streamlit as st
import google.generativeai as genai
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from datetime import datetime
import uuid
import io
from docx import Document

# --- 1. KONFIGURACE A STYLY ---
st.set_page_config(page_title="S.M.A.R.T. OS", page_icon="🤖", layout="wide")

st.markdown("""
    <style>
    [data-testid="stStatusWidget"], .stDeployButton, footer, #MainMenu, header { display: none !important; }
    .thinking-box {
        background-color: #1e2129; border-left: 5px solid #0288d1;
        padding: 15px; border-radius: 5px; color: #e0e0e0;
        font-weight: bold; margin: 10px 0;
    }
    .fixed-footer {
        position: fixed; left: 0; bottom: 0; width: 100%;
        text-align: center; color: #888; font-size: 0.8rem;
        padding: 15px; background: #0e1117;
        border-top: 1px solid #262730; z-index: 1000;
    }
    .main-content { margin-bottom: 120px; }
    </style>
    """, unsafe_allow_html=True)

# --- 2. EXPORT WORD ---
def create_docx(text):
    doc = Document()
    doc.add_heading('S.M.A.R.T. OS Výstup', 0)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. SESSION STATE ---
if "user_id" not in st.session_state: st.session_state.user_id = str(uuid.uuid4())[:8]
if "chat_id" not in st.session_state: st.session_state.chat_id = str(uuid.uuid4())[:8]
if "last_response" not in st.session_state: st.session_state.last_response = ""

# --- 4. DATABÁZE ---
conn = st.connection("gsheets", type=GSheetsConnection)

def load_data():
    try:
        users = conn.read(worksheet="Users", ttl=0)
        stats = conn.read(worksheet="Stats", ttl=0)
        return users, stats
    except:
        return pd.DataFrame(columns=["user_id", "chat_id", "title", "role", "content", "timestamp"]), \
               pd.DataFrame(columns=["key", "value"])

users_df, stats_df = load_data()

try:
    total_msgs = int(stats_df[stats_df['key'] == 'total_messages']['value'].values[0])
except:
    total_msgs = 0

# --- 5. SIDEBAR ---
with st.sidebar:
    st.title("🤖 S.M.A.R.T. OS")
    # Dynamický název modelu pro informaci uživatele
    active_label = "Gemini 3" if total_msgs < 200 else "Gemini 2.5 Lite"
    st.caption(f"🚀 Režim: {active_label}")
    
    if st.button("➕ Nový chat", use_container_width=True):
        st.session_state.chat_id = str(uuid.uuid4())[:8]
        st.session_state.last_response = ""
        st.rerun()
    
    st.divider()
    uploaded_file = st.file_uploader("Nahrát podklady", type=["png", "jpg", "jpeg", "pdf", "txt"])
    
    if st.session_state.last_response:
        st.download_button("📥 Stáhnout jako Word", data=create_docx(st.session_state.last_response), file_name="smart_export.docx", use_container_width=True)

# --- 6. CHAT ---
current_chat = users_df[users_df["chat_id"] == st.session_state.chat_id]
st.subheader("💬 S.M.A.R.T. Chat")

st.markdown('<div class="main-content">', unsafe_allow_html=True)
for _, m in current_chat.iterrows():
    with st.chat_message(m["role"]): st.write(m["content"])

# --- 7. LOGIKA (POUZE 3 A 2.5 FLASH LITE) ---
if prompt := st.chat_input("Zeptejte se na cokoliv..."):
    with st.chat_message("user"): st.write(prompt)
    
    thinking = st.empty()
    thinking.markdown('<div class="thinking-box">🤖 SMART přemýšlí...</div>', unsafe_allow_html=True)
    
    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    api_keys = [st.secrets.get(f"GOOGLE_API_KEY_{i}") for i in range(1, 11)]
    
    # Prioritní model podle limitu 200 zpráv
    primary_model = "gemini-3-flash" if total_msgs < 200 else "gemini-2.5-flash-lite"
    # Záložní model (vždy 2.5 lite)
    fallback_model = "gemini-2.5-flash-lite"
    
    success = False
    ai_text = ""

    history = []
    for _, m in current_chat.tail(10).iterrows():
        history.append({"role": "user" if m["role"] == "user" else "model", "parts": [m["content"]]})

    parts = [prompt]
    if uploaded_file:
        f_bytes = uploaded_file.read()
        if uploaded_file.type == "text/plain":
            parts.append(f"\n\nSoubor: {uploaded_file.name}\n{f_bytes.decode('utf-8')}")
        else:
            parts.append({"mime_type": uploaded_file.type, "data": f_bytes})

    # ROTACE KLÍČŮ A MODELŮ
    for key in api_keys:
        if not key or success: continue
        try:
            genai.configure(api_key=key)
            
            # Pokus o vygenerování s primárním modelem
            try:
                # Zkusíme s vyhledáváním
                model = genai.GenerativeModel(
                    model_name=primary_model,
                    tools=[{"google_search_retrieval": {}}],
                    system_instruction="Jsi S.M.A.R.T. OS. Pomáhej studentům. Pro matiku LaTeX ($$). Odpovídej česky."
                )
                chat = model.start_chat(history=history)
                response = chat.send_message(parts)
                ai_text = response.text
                success = True
            except:
                # Pokud vyhledávání nebo v3 selže, zkusíme 2.5 Lite bez vyhledávání
                model = genai.GenerativeModel(model_name=fallback_model)
                chat = model.start_chat(history=history)
                response = chat.send_message(parts)
                ai_text = response.text
                success = True
        except:
            continue

    thinking.empty()

    if success:
        with st.chat_message("assistant"): st.markdown(ai_text)
        st.session_state.last_response = ai_text
        
        # Uložení do GSheets
        u_row = pd.DataFrame([{"user_id": st.session_state.user_id, "chat_id": st.session_state.chat_id, "title": prompt[:20], "role": "user", "content": prompt, "timestamp": now}])
        ai_row = pd.DataFrame([{"user_id": st.session_state.user_id, "chat_id": st.session_state.chat_id, "title": prompt[:20], "role": "assistant", "content": ai_text, "timestamp": now}])
        conn.update(worksheet="Users", data=pd.concat([users_df, u_row, ai_row], ignore_index=True))
        
        # Update počítadla
        stats_df.loc[stats_df['key'] == 'total_messages', 'value'] = str(total_msgs + 1)
        conn.update(worksheet="Stats", data=stats_df)
        
        st.rerun()
    else:
        st.error("❌ Systém je přetížen. Zkuste to za chvíli.")

st.markdown('</div>', unsafe_allow_html=True)
st.markdown('<div class="fixed-footer">S.M.A.R.T. OS může dělat chyby.</div>', unsafe_allow_html=True)
